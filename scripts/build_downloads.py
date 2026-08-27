#!/usr/bin/env python3
"""Build styled downloadable CV assets to match the GitHub Pages CV site."""
from __future__ import annotations

import json
import math
import re
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph
from reportlab.pdfgen import canvas

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
DATA = json.loads((ASSETS / "cv-content.json").read_text(encoding="utf-8"))
PDF_OUT = ASSETS / "William_McLaughlan_CV.pdf"
DOCX_OUT = ASSETS / "William_McLaughlan_CV.docx"
MD_OUT = ASSETS / "William_McLaughlan_CV.md"
HEADSHOT = ASSETS / "headshot.png"

PAGE_W, PAGE_H = A4
SIDE_W = 73 * mm
MARGIN = 13 * mm
MAIN_X = SIDE_W + 13 * mm
MAIN_W = PAGE_W - MAIN_X - MARGIN

COLORS = {
    "bg": colors.HexColor("#080b08"),
    "side": colors.HexColor("#243328"),
    "side2": colors.HexColor("#172119"),
    "side_line": colors.Color(1, 1, 1, alpha=0.18),
    "side_text": colors.HexColor("#f4f6ef"),
    "side_muted": colors.HexColor("#c5cfbd"),
    "accent": colors.HexColor("#c9c1a9"),
    "accent2": colors.HexColor("#a3b18a"),
    "paper": colors.HexColor("#f6f5ef"),
    "ink": colors.HexColor("#151813"),
    "muted": colors.HexColor("#595d55"),
    "rule": colors.HexColor("#d6d1c2"),
    "white_soft": colors.HexColor("#eef1ea"),
}

BODY = ParagraphStyle("body", fontName="Helvetica", fontSize=7.7, leading=9.45, textColor=COLORS["muted"], spaceAfter=3)
BULLET = ParagraphStyle("bullet", fontName="Helvetica", fontSize=7.0, leading=8.5, textColor=colors.HexColor("#353a33"), leftIndent=9, firstLineIndent=-7, spaceAfter=2.0)
TITLE = ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=10.5, leading=12.5, textColor=COLORS["ink"], spaceAfter=3)
ROLE = ParagraphStyle("role", fontName="Helvetica-Bold", fontSize=8.0, leading=9.8, textColor=COLORS["ink"], spaceAfter=1)
SUB = ParagraphStyle("sub", fontName="Helvetica", fontSize=6.8, leading=8.2, textColor=COLORS["muted"], spaceAfter=4)
SIDE_TITLE = ParagraphStyle("side_title", fontName="Helvetica-Bold", fontSize=9.2, leading=11, textColor=COLORS["side_text"], spaceAfter=6)
SIDE_TEXT = ParagraphStyle("side_text", fontName="Helvetica", fontSize=7.2, leading=8.7, textColor=COLORS["side_text"], spaceAfter=3)
SIDE_MUTED = ParagraphStyle("side_muted", fontName="Helvetica", fontSize=6.8, leading=8.2, textColor=COLORS["side_muted"], spaceAfter=4)

TECH = [
    ("M365", "Microsoft 365"), ("AZ", "Azure"), ("WIN", "Windows"),
    ("AD", "Active Directory"), ("INT", "Intune"), ("JAMF", "JAMF"),
    ("ZOOM", "Zoom"), ("D365", "Dynamics 365"), ("RM", "Roundsman"),
    ("TEAMS", "Teams"), ("CTX", "Citrix"), ("MIME", "Mimecast"),
    ("EX", "Exchange"), ("RDP", "Remote"), ("TV", "TeamViewer"),
    ("FORT", "Fortinet"), ("ITIL", "ITIL v4"),
]
TECH_ICONS = {
    "Microsoft 365": "Microsoft365.png",
    "Azure": "Microsoft-Azure.png",
    "Windows": "Microsoft-WindowsLogoColour.png",
    "Active Directory": "Microsoft-Azure.png",
    "Intune": "Microsoft-IntuneCompanyPortal.png",
    "Zoom": "Zoom.png",
    "Dynamics 365": "Microsoft-Dynamics365.png",
    "Teams": "Microsoft-Teams.png",
    "Citrix": "Citrix-WorkspaceApp.png",
    "Mimecast": "Mimecast.png",
    "Exchange": "Microsoft-Outlook.png",
    "Remote": "Microsoft-RemoteDesktop.png",
    "Defender": "Microsoft-Defender.png",
}
BRAND_MARKS = {
    "JAMF": "jamf",
    "Roundsman": "RM",
    "TeamViewer": "TV",
    "Fortinet": "FORTINET",
    "ITIL v4": "ITIL",
}
CHIPS = [
    "Service delivery", "Infrastructure", "Cyber security", "Microsoft 365 admin",
    "Entra ID / AD", "GPOs", "MDM: Intune / SCCM / JAMF", "Hybrid Exchange",
    "SaaS support", "Bespoke software", "Hardware support", "Imaging / FOG / Sysprep",
    "Asset management", "Knowledge base creation", "VoIP: RingCentral / Avaya",
    "VPN: FortiClient", "Citrix Workspace", "Suppliers / MSPs", "Licensing",
    "Projects", "Stakeholders", "Budget-ready",
]


def esc(text: str) -> str:
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def para(c: canvas.Canvas, text: str, style: ParagraphStyle, x: float, y: float, w: float, max_h: float = 1000) -> float:
    p = Paragraph(text, style)
    _, h = p.wrap(w, max_h)
    p.drawOn(c, x, y - h)
    return y - h - style.spaceAfter


def round_rect(c, x, y, w, h, radius, fill, stroke=None, lw=1.0):
    c.setFillColor(fill)
    if stroke:
        c.setStrokeColor(stroke); c.setLineWidth(lw)
        c.roundRect(x, y, w, h, radius, fill=1, stroke=1)
    else:
        c.roundRect(x, y, w, h, radius, fill=1, stroke=0)


def draw_page_bg(c: canvas.Canvas):
    c.setFillColor(COLORS["paper"])
    c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    strips = 70
    start = (36, 51, 40)
    end = (23, 33, 25)
    for i in range(strips):
        t = i / (strips - 1)
        rgb = tuple((start[j] * (1 - t) + end[j] * t) / 255 for j in range(3))
        c.setFillColor(colors.Color(*rgb))
        c.rect(0, PAGE_H - (i + 1) * PAGE_H / strips, SIDE_W, PAGE_H / strips + 1, fill=1, stroke=0)
    c.setStrokeColor(colors.Color(1, 1, 1, alpha=0.12))
    c.line(SIDE_W, 0, SIDE_W, PAGE_H)


def draw_headshot(c, cx, cy, size):
    if HEADSHOT.exists():
        c.saveState()
        p = c.beginPath()
        p.circle(cx, cy, size / 2)
        c.clipPath(p, stroke=0, fill=0)
        c.drawImage(str(HEADSHOT), cx - size / 2, cy - size / 2, width=size, height=size, preserveAspectRatio=True, anchor="c")
        c.restoreState()
    c.setStrokeColor(colors.Color(1, 1, 1, alpha=0.35))
    c.setLineWidth(2)
    c.circle(cx, cy, size / 2, fill=0, stroke=1)


def side_rule(c, y):
    c.setStrokeColor(colors.Color(1, 1, 1, alpha=0.18))
    c.line(12 * mm, y, SIDE_W - 12 * mm, y)


def draw_tech_logo(c: canvas.Canvas, label: str, cell_x: float, cell_y: float, cell_w: float, cell_h: float):
    """Draw product logo/mark only — no surrounding card or inner box."""
    cx = cell_x + cell_w / 2
    icon_cy = cell_y + cell_h * 0.58
    icon = TECH_ICONS.get(label)
    if icon and (ASSETS / "icons" / icon).exists():
        size = 7.8 * mm
        c.drawImage(str(ASSETS / "icons" / icon), cx - size / 2, icon_cy - size / 2, width=size, height=size, preserveAspectRatio=True, anchor="c", mask="auto")
        return
    mark = BRAND_MARKS.get(label, label[:4].upper())
    c.setFillColor(COLORS["side_text"])
    c.setFont("Helvetica-Bold", 3.9 if len(mark) > 4 else 5.0)
    c.drawCentredString(cx, icon_cy - 1.5 * mm, mark)


def draw_sidebar(c: canvas.Canvas, page_num: int):
    x = 12 * mm
    w = SIDE_W - 24 * mm
    y = PAGE_H - 16 * mm
    if page_num == 1:
        draw_headshot(c, SIDE_W / 2, PAGE_H - 34 * mm, 32 * mm)
        y = PAGE_H - 58 * mm
    if page_num == 1:
        side_rule(c, y); y -= 7 * mm
        y = para(c, "CONTACT", SIDE_TITLE, x, y, w)
        contact = DATA["contact"]
        contacts = [contact["email"], contact["phone"], "Glasgow, Scotland", "linkedin.com/in/william-mclaughlan"]
        for item in contacts:
            y = para(c, esc(item), SIDE_TEXT, x, y, w)
        side_rule(c, y - 2 * mm); y -= 9 * mm
        y = para(c, "TECHNOLOGY", SIDE_TITLE, x, y, w)
        cell_w = (w - 5 * mm) / 3
        cell_h = 13.8 * mm
        for idx, (abbr, label) in enumerate(TECH):
            col = idx % 3
            row = idx // 3
            tx = x + col * (cell_w + 2.5 * mm)
            ty = y - (row + 1) * cell_h
            draw_tech_logo(c, label, tx, ty, cell_w, cell_h)
            c.setFillColor(COLORS["side_muted"]); c.setFont("Helvetica", 4.8)
            c.drawCentredString(tx + cell_w / 2, ty + 3.0 * mm, label[:17])
        y -= math.ceil(len(TECH) / 3) * cell_h + 4 * mm
        side_rule(c, y); y -= 7 * mm
        y = para(c, "SKILLS", SIDE_TITLE, x, y, w)
        y = draw_chips(c, CHIPS[:12], x, y, w)
        y -= 3 * mm
        para(c, "Continued on next page", SIDE_MUTED, x, y, w)
    elif page_num == 2:
        side_rule(c, y); y -= 7 * mm
        y = para(c, "SKILLS CONTINUED", SIDE_TITLE, x, y, w)
        y = draw_chips(c, CHIPS[12:], x, y, w)
        y -= 8 * mm
        side_rule(c, y); y -= 7 * mm
        y = para(c, "OUTSIDE OF IT", SIDE_TITLE, x, y, w)
        para(c, "Game development • Surfing • Camping • Reading", SIDE_MUTED, x, y, w)
    elif page_num == 3:
        # Intentional blank branded sidebar: do not repeat contact/core-strength
        # content on later pages.
        return


def draw_chips(c: canvas.Canvas, chips: list[str], x: float, y: float, w: float) -> float:
    chip_y = y
    chip_x = x
    for chip in chips:
        text_w = c.stringWidth(chip, "Helvetica", 5.5) + 5.4 * mm
        if chip_x + text_w > x + w:
            chip_x = x
            chip_y -= 6.6 * mm
        round_rect(c, chip_x, chip_y - 4.3 * mm, text_w, 5.3 * mm, 2.6 * mm, colors.Color(1, 1, 1, alpha=0.10), colors.Color(1, 1, 1, alpha=0.14), 0.4)
        c.setFillColor(COLORS["side_text"]); c.setFont("Helvetica", 5.5)
        c.drawString(chip_x + 2.2 * mm, chip_y - 2.6 * mm, chip)
        chip_x += text_w + 1.5 * mm
    return chip_y - 6.5 * mm


def draw_hero(c, y):
    x = MAIN_X
    h = 33 * mm
    round_rect(c, x, y - h, MAIN_W, h, 2.5 * mm, colors.HexColor("#ebe8dd"), colors.HexColor("#ddd6c4"), 0.7)
    c.setFillColor(COLORS["ink"])
    c.setFont("Helvetica-Bold", 19.5)
    c.drawCentredString(x + MAIN_W / 2, y - 13 * mm, DATA["name"].upper())
    c.setStrokeColor(colors.HexColor("#777366")); c.setLineWidth(1.1)
    c.line(x + MAIN_W / 2 - 34 * mm, y - 17.5 * mm, x + MAIN_W / 2 + 34 * mm, y - 17.5 * mm)
    c.setFillColor(COLORS["muted"]); c.setFont("Helvetica-Bold", 7.7)
    c.drawCentredString(x + MAIN_W / 2, y - 24 * mm, "IT SERVICE DESK MANAGER • INFRASTRUCTURE • SERVICE DELIVERY")
    return y - h - 7 * mm


def draw_section_title(c, title, y):
    c.setFillColor(COLORS["ink"]); c.setFont("Helvetica-Bold", 12)
    c.drawString(MAIN_X, y, title)
    c.setStrokeColor(COLORS["rule"]); c.setLineWidth(0.8)
    c.line(MAIN_X + c.stringWidth(title, "Helvetica-Bold", 12) + 6 * mm, y + 2, MAIN_X + MAIN_W, y + 2)
    return y - 6 * mm


def fit_para_height(text, style, w):
    p = Paragraph(text, style)
    _, h = p.wrap(w, 1000)
    return h + style.spaceAfter


def split_earlier_career(item: str):
    parts = [p.strip() for p in item.split(" — ")]
    if len(parts) >= 3:
        return parts[0], " — ".join(parts[1:-1]), parts[-1]
    if len(parts) == 2:
        return parts[0], "", parts[1]
    return item, "", ""


def draw_earlier_career_grid(c: canvas.Canvas, items: list[str], x: float, y: float, w: float) -> float:
    cols = 2
    gap = 4 * mm
    card_w = (w - gap) / cols
    card_h = 12.8 * mm
    for idx, item in enumerate(items):
        col = idx % cols
        row = idx // cols
        cx = x + col * (card_w + gap)
        cy = y - row * (card_h + 2.2 * mm) - card_h
        round_rect(c, cx, cy, card_w, card_h, 2 * mm, colors.HexColor("#f0ede2"), colors.HexColor("#ded7c7"), 0.45)
        company, role, dates = split_earlier_career(item)
        c.setFillColor(COLORS["ink"]); c.setFont("Helvetica-Bold", 6.2)
        c.drawString(cx + 2.6 * mm, cy + 8.3 * mm, company[:31])
        c.setFillColor(COLORS["muted"]); c.setFont("Helvetica", 5.7)
        role_text = role[:34] if role else "Earlier education / transition"
        c.drawString(cx + 2.6 * mm, cy + 5.0 * mm, role_text)
        c.setFillColor(colors.HexColor("#777366")); c.setFont("Helvetica-Bold", 5.2)
        c.drawString(cx + 2.6 * mm, cy + 2.1 * mm, dates[:33])
    rows = math.ceil(len(items) / cols)
    return y - rows * (card_h + 2.2 * mm) - 2 * mm


def draw_page_three_footer(c: canvas.Canvas):
    panel_x = MAIN_X
    panel_y = 22 * mm
    panel_h = 32 * mm
    round_rect(c, panel_x, panel_y, MAIN_W, panel_h, 2.5 * mm, colors.HexColor("#ebe8dd"), colors.HexColor("#ddd6c4"), 0.6)
    c.setFillColor(COLORS["ink"]); c.setFont("Helvetica-Bold", 9.5)
    c.drawString(panel_x + 5 * mm, panel_y + panel_h - 8 * mm, "SERVICE DELIVERY • SYSTEMS • CHANGE")
    c.setFillColor(COLORS["muted"]); c.setFont("Helvetica", 6.5)
    c.drawString(panel_x + 5 * mm, panel_y + panel_h - 15 * mm, "Practical IT leadership across BAU support, infrastructure, business systems and supplier coordination.")
    c.drawString(panel_x + 5 * mm, panel_y + panel_h - 21 * mm, "Hands-on delivery with clear stakeholder communication and steady technology change.")


def ensure_space(c, y, needed, page_num):
    if y - needed < 18 * mm:
        c.showPage()
        draw_page_bg(c); draw_sidebar(c, 2)
        return PAGE_H - 18 * mm, page_num + 1
    return y, page_num


def draw_experience_item(c, exp, y, page_num, max_bullets=None):
    bullets = exp["bullets"] if max_bullets is None else exp["bullets"][:max_bullets]
    y, page_num = ensure_space(c, y, 18 * mm, page_num)
    header = f"<b>{esc(exp['role'])} — {esc(exp['company'])}</b>"
    y = para(c, header, ROLE, MAIN_X, y, MAIN_W)
    y = para(c, f"<b>{esc(exp['dates'])}</b>", SUB, MAIN_X, y, MAIN_W)
    y = para(c, esc(exp.get("intro", "")), SUB, MAIN_X, y, MAIN_W)
    for b in bullets:
        needed = fit_para_height("• " + esc(b), BULLET, MAIN_W - 6 * mm)
        y, page_num = ensure_space(c, y, needed + 2, page_num)
        y = para(c, "• " + esc(b), BULLET, MAIN_X + 4 * mm, y, MAIN_W - 4 * mm)
    c.setStrokeColor(colors.HexColor("#dfd8c8")); c.setLineWidth(0.5)
    c.line(MAIN_X, y - 0.8 * mm, MAIN_X + MAIN_W, y - 0.8 * mm)
    return y - 4 * mm, page_num


def build_pdf():
    c = canvas.Canvas(str(PDF_OUT), pagesize=A4)
    c.setTitle("William McLaughlan — CV")
    draw_page_bg(c); draw_sidebar(c, 1)
    y = PAGE_H - 16 * mm
    y = draw_hero(c, y)
    y = draw_section_title(c, "SUMMARY", y)
    profile_parts = re.split(r"(?<=depends on\. )", DATA["profile"], maxsplit=1)
    for part in profile_parts:
        if part.strip():
            y = para(c, esc(part.strip()), BODY, MAIN_X, y, MAIN_W)
    y -= 2 * mm
    # Keep the downloadable PDF polished and readable rather than cramming the whole
    # web page into two pages. Explicit breaks avoid the old-style CV's crowding and
    # keep the hosted-page visual language intact across pages.
    y -= 5 * mm
    y = draw_section_title(c, "EXPERIENCE", y)
    page_num = 1
    y, page_num = draw_experience_item(c, DATA["experience"][0], y, page_num)

    c.showPage()
    page_num += 1
    draw_page_bg(c); draw_sidebar(c, 2)
    y = PAGE_H - 18 * mm
    y = draw_section_title(c, "EXPERIENCE CONTINUED", y)
    y, page_num = draw_experience_item(c, DATA["experience"][1], y, page_num)
    y, page_num = draw_experience_item(c, DATA["experience"][2], y, page_num)

    c.showPage()
    page_num += 1
    draw_page_bg(c); draw_sidebar(c, 3)
    y = PAGE_H - 18 * mm
    y = draw_section_title(c, "EXPERIENCE CONTINUED", y)
    y, page_num = draw_experience_item(c, DATA["experience"][3], y, page_num)
    y -= 7 * mm
    earlier_needed = 66 * mm
    y, page_num = ensure_space(c, y, earlier_needed, page_num)
    y = draw_section_title(c, "EARLIER CAREER", y)
    y = draw_earlier_career_grid(c, DATA["earlier_career"], MAIN_X, y, MAIN_W)
    y -= 5 * mm
    y = draw_section_title(c, "CERTIFICATION", y)
    para(c, esc(", ".join(DATA["certifications"])), BODY, MAIN_X, y, MAIN_W)
    draw_page_three_footer(c)
    c.save()


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="FFFFFF", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_cell_margins(table, top=120, left=180, bottom=120, right=180):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for edge, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = tbl_cell_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=120, left=220, bottom=120, right=220):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cell_para(cell, text, size=8.5, bold=False, italic=False, color="F4F6EF", align=None, before=0, after=3):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def build_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
    sec.top_margin = Inches(0); sec.bottom_margin = Inches(0)
    sec.left_margin = Inches(0); sec.right_margin = Inches(0)
    sec.header_distance = Inches(0); sec.footer_distance = Inches(0)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_cell_margins(table, top=120, left=180, bottom=120, right=180)
    table.rows[0].height = Inches(11.69)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    side, main = table.rows[0].cells
    side.width = Inches(2.55); main.width = Inches(5.72)
    set_cell_margins(side, top=180, left=220, bottom=180, right=180)
    set_cell_margins(main, top=180, left=280, bottom=180, right=300)
    set_cell_shading(side, "243328"); set_cell_shading(main, "F6F5EF")
    set_cell_border(side, "243328"); set_cell_border(main, "F6F5EF")
    side.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    main.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if HEADSHOT.exists():
        p = side.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(HEADSHOT), width=Inches(1.25))
    cell_para(side, "CONTACT", 12, True, color="F4F6EF", before=8, after=5)
    contact = DATA["contact"]
    for item in [contact["email"], contact["phone"], "Glasgow, Scotland", "linkedin.com/in/william-mclaughlan"]:
        cell_para(side, item, 7.4, color="F4F6EF", after=2)
    cell_para(side, "TECHNOLOGY", 12, True, color="F4F6EF", before=8, after=5)
    for i in range(0, len(TECH), 2):
        labels = "   ".join(label for _, label in TECH[i:i+2])
        cell_para(side, labels, 6.3, color="C5CFBD", after=1)
    cell_para(side, "SKILLS", 12, True, color="F4F6EF", before=8, after=5)
    for chip in CHIPS:
        cell_para(side, "  " + chip + "  ", 6.7, color="F4F6EF", after=1)
    cell_para(side, "OUTSIDE OF IT", 12, True, color="F4F6EF", before=8, after=4)
    cell_para(side, "Game development • Surfing • Camping • Reading", 7.0, color="C5CFBD", after=2)

    p = main.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(DATA["name"].upper())
    r.bold = True; r.font.name = "Georgia"; r.font.size = Pt(22); r.font.color.rgb = RGBColor.from_string("151813")
    p.paragraph_format.space_after = Pt(2)
    p = main.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IT SERVICE DESK MANAGER • INFRASTRUCTURE • SERVICE DELIVERY")
    r.bold = True; r.font.size = Pt(7.5); r.font.color.rgb = RGBColor.from_string("595D55")
    p.paragraph_format.space_after = Pt(9)

    def h(text):
        cell_para(main, text, 12, True, color="151813", before=5, after=4)
    def mpara(text, size=8.2, italic=False):
        cell_para(main, text, size, italic=italic, color="595D55", after=4)
    def bullet(text):
        p = main.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(text)
        r.font.name = "Aptos"; r.font.size = Pt(7.4); r.font.color.rgb = RGBColor.from_string("353A33")

    h("SUMMARY")
    for part in re.split(r"(?<=business\. )", DATA["profile"], maxsplit=1):
        if part.strip():
            mpara(part.strip(), 8.0)
    h("EXPERIENCE")
    for exp in DATA["experience"]:
        cell_para(main, f"{exp['role']} — {exp['company']}    {exp['dates']}", 8.8, True, color="151813", after=1)
        mpara(exp["intro"], 7.6, True)
        for b in exp["bullets"]:
            bullet(b)
    h("EARLIER CAREER")
    for item in DATA["earlier_career"]:
        company, role, dates = split_earlier_career(item)
        p = main.add_paragraph()
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(company)
        r.bold = True; r.font.name = "Aptos"; r.font.size = Pt(7.6); r.font.color.rgb = RGBColor.from_string("151813")
        detail = f" — {role} — {dates}" if role else f" — {dates}"
        r = p.add_run(detail)
        r.font.name = "Aptos"; r.font.size = Pt(7.2); r.font.color.rgb = RGBColor.from_string("595D55")
    h("CERTIFICATION")
    mpara(", ".join(DATA["certifications"]), 7.8)
    p = main.add_paragraph()
    p.paragraph_format.space_after = Pt(220)
    p = side.add_paragraph()
    p.paragraph_format.space_after = Pt(220)
    doc.save(DOCX_OUT)


def build_md():
    c = DATA["contact"]
    lines = [
        f"# {DATA['name']}",
        " | ".join([c["email"], c["phone"], c["linkedin"], c["location"]]),
        "", "## Summary", DATA["profile"], "", "## Skills",
    ]
    lines += [f"- {s}" for s in DATA["core_skills"]]
    lines += ["", "## Experience"]
    for exp in DATA["experience"]:
        lines += [f"### {exp['role']} — {exp['company']} — {exp['dates']}", exp["intro"]]
        lines += [f"- {b}" for b in exp["bullets"]]
    lines += ["", "## Earlier Career"] + [f"- {e}" for e in DATA["earlier_career"]]
    lines += ["", "## Certifications"] + [f"- {cert}" for cert in DATA["certifications"]]
    MD_OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    build_pdf()
    build_docx()
    build_md()
    print(PDF_OUT)
    print(DOCX_OUT)
    print(MD_OUT)
